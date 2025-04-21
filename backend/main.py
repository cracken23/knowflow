import sys
import os
from io import BytesIO
import tempfile
import requests
from git import Repo
from pathlib import Path
from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document

from flow import create_paper_flow

app = Flask(__name__)
CORS(app)
# Both endpoints have identical DOCX creation logic - extract to a function
def generate_docx_from_sections(paper_sections: dict) -> BytesIO:
    """Shared DOCX generation logic for both endpoints."""
    doc = Document()
    doc.add_heading(paper_sections.get("Title", "Untitled IEEE Paper"), level=0)
    
    for section in [
        "Abstract", "Introduction", "Methodology",
        "Results & Discussion", "Conclusion", "References"
    ]:
        if content := paper_sections.get(section):
            doc.add_heading(section, level=1)
            for para in content.split("\n\n"):
                doc.add_paragraph(para.strip())
    
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
def generate_documentation_from_code(code: str) -> str:
    # print(code)
    if not os.getenv('OPENROUTER_API_KEY'):
        raise ValueError("OPENROUTER_API_KEY environment variable missing")
    """Use LLM to generate documentation for a code snippet."""
    try:
        url = "https://openrouter.ai/api/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY')}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "qwen/qwen-2.5-coder-32b-instruct:free",  # example open‑source model 
            "messages": [{"role": "user", "content": f"Code to document:\n{code}"}],
            "temperature": 0.3,
            "stream": False
        }

        resp = requests.post(url, json=payload, headers=headers)
        resp.raise_for_status()
        data = resp.json()
        return data["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"Error generating docs: {e}")
        return f"Documentation generation failed: {str(e)}"
    
def process_github_repo(repo_url: str) -> str:
    """Clone repo and process all code files."""
    docs = []
    with tempfile.TemporaryDirectory() as tmp_dir:
        try:
            # Add branch/tag validation if needed
            Repo.clone_from(repo_url, tmp_dir, depth=1)
            
            # Skip binary files
            text_chars = bytearray({7,8,9,10,12,13,27} | set(range(0x20, 0x100)) - {0x7f})
            def is_binary_file(path: Path):
                return bool(open(path, 'rb').read(1024).translate(None, text_chars))

            code_extensions = {'.py', '.js', '.java', '.c', '.cpp', '.h', '.ts'}
            for path in Path(tmp_dir).rglob('*'):
                if path.suffix in code_extensions and path.is_file() and not is_binary_file(path):
                    try:
                        with open(path, 'r', encoding='utf-8') as f:
                            code = f.read()
                            # Truncate large files
                            if len(code) > 100000:  # 100KB limit per file
                                code = code[:100000] + "\n// ... [truncated due to size] ..."
                            file_doc = f"=== File: {path.relative_to(tmp_dir)} ===\n"
                            file_doc += generate_documentation_from_code(code)
                            docs.append(file_doc)
                    except UnicodeDecodeError:
                        print(f"Skipping non-text file: {path}")
                    except Exception as e:
                        print(f"Error processing {path}: {e}")
            
            return "\n\n".join(docs)
        except Exception as e:
            print(f"Repo processing error: {str(e)}")
            raise
@app.route('/api/generate_paper', methods=['POST'])
def generate_paper():
    try:
        data = request.get_json()
        if not (doc_text := data.get("documentation")):
            return {"error": "Missing documentation payload"}, 400
            
        agent_flow = create_paper_flow()
        shared = {"documentation": doc_text}
        agent_flow.run(shared)
        
        if not (paper_sections := shared.get("paper_sections")):
            return {"error": "Paper generation failed - no sections produced"}, 500
            
        return send_file(
            generate_docx_from_sections(paper_sections),
            as_attachment=True,
            download_name='ieee_paper.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"General paper generation error: {str(e)}")
        return {"error": f"Paper generation failed: {str(e)}"}, 500

@app.route('/api/generate_from_github', methods=['POST'])
def generate_from_github():
    print("Received request to generate paper from GitHub repo.")
    data = request.get_json()
    print(f"Request data: {data}")
    repo_url = data.get("repo_url")
    
    if not repo_url:
        app.logger.error("GitHub URL not provided.")
        return {"error": "GitHub URL required"}, 400
    
    if not repo_url.startswith('https://github.com/'):
        return {"error": "Invalid GitHub URL format"}, 400

    try:
        app.logger.info(f"Cloning repository from {repo_url}")
        generated_docs = process_github_repo(repo_url)
        agent_flow = create_paper_flow()
        shared = {"documentation": generated_docs}
        agent_flow.run(shared)
        
        if not (paper_sections := shared.get("paper_sections")):
            return {"error": "Paper generation failed - no sections produced"}, 500
            
        return send_file(
            generate_docx_from_sections(paper_sections),
            as_attachment=True,
            download_name='ieee_paper.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error processing GitHub repo: {str(e)}")
        return {"error": f"Failed to process repository: {str(e)}"}, 500
    
def run_cli():
    """Fallback CLI mode—unchanged."""
    default_documentation = """
    [Insert your doxygen-generated code documentation here]
    ...
    """
    doc_input = default_documentation
    if len(sys.argv) > 1 and sys.argv[1] != "serve":
        try:
            with open(sys.argv[1], 'r', encoding='utf-8') as f:
                doc_input = f.read()
        except Exception as e:
            print(f"Error reading file: {e}")
            sys.exit(1)

    agent_flow = create_paper_flow()
    shared = {"documentation": doc_input}
    agent_flow.run(shared)
    print(shared.get("final_paper", "No paper generated"))

if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "serve":
        app.run(port=5000, debug=True)
    else:
        run_cli()
